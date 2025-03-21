import tkinter as tk
from tkinter import ttk
from enum import Enum
from tkinter import messagebox
import docx
import openpyxl
from openpyxl.utils.cell import get_column_letter

# Model

# Границы передвижения тележек с поста на пост
MIN_TRANSFER_TIME = 1.5
MAX_TRANSFER_TIME = 3

class OperationType(Enum):
    AUTOMATED = "Автоматизированные"
    MECHANIZED = "Механизированные"
    MANUAL = "Ручные"

class ProductComplexity(Enum):
    SIMPLE = "Изделия простой конфигурации, однослойные"
    COMPLEX = "Изделия сложной конфигурации, многослойные или офактуренные"

"""Сервис для полцучения коэффициентов конвейерной линии"""
class OperationCoefficients:
    def __init__(self, operation_type, product_complexity, coefficient_numerator, coefficient_denominator=None):
        
        self.operation_type = operation_type
        self.product_complexity = product_complexity
        self.coefficient_numerator = coefficient_numerator
        self.coefficient_denominator = coefficient_denominator

    def get_coefficient(self, use_denominator=False):
        return self.coefficient_denominator if use_denominator and self.coefficient_denominator is not None else self.coefficient_numerator
    
class ConveyorLineSegment:
    def __init__(self, operation_type, product_complexity, preform_operation_time, form_operation_time, postform_operation_time, cycle_time, transfer_time):

        if operation_type not in OperationType:
            raise ValueError("Операция не определена!")
        if product_complexity not in ProductComplexity:
            raise ValueError("Сложность изделия не определена!")
        if preform_operation_time < 0:
            raise ValueError("Средняя продолжительность операции на доформочном участке не может быть отрицательной!")
        if form_operation_time < 0:
            raise ValueError("Средняя продолжительность операции на формочном участке не может быть отрицательной!")
        if postform_operation_time < 0:
            raise ValueError("Средняя продолжительность операции на постформочном участке не может быть отрицательной!")
        if cycle_time < 0:
            raise ValueError("Продолжительность цикла формования не может быть отрицательной!")
        if cycle_time <= transfer_time:
            raise ValueError("Продолжительность цикла формования должна быть больше продолжительности передвижения тележек!")
        if MIN_TRANSFER_TIME > transfer_time or transfer_time > MAX_TRANSFER_TIME:
            raise ValueError(f'Продолжительность передвижения тележек должна быть в диапазоне от {MIN_TRANSFER_TIME} до {MAX_TRANSFER_TIME}!')

        self.operation_type = operation_type
        self.product_complexity = product_complexity
        self.preform_operation_time = preform_operation_time
        self.form_operation_time = form_operation_time
        self.postform_operation_time = postform_operation_time
        self.cycle_time = cycle_time
        self.transfer_time = transfer_time

""" Сервис для расчета постов конвейерной линии """
class Calculator:
    def calculate_posts(self, segment, coefficient, use_denominator=False):
        kh = coefficient.get_coefficient(use_denominator)
        result = ((segment.preform_operation_time + segment.form_operation_time + segment.postform_operation_time) * kh) / (segment.cycle_time - segment.transfer_time)
        if result >= 0:
            return result
        else:
            raise ValueError("Результат расчета не может быть отрицательным!")

""" Сервис для экспорта данных в файлы"""
class ExportService:
    def export_to_docx(self, filename, data):
        doc = docx.Document()
        doc.add_heading("Результаты расчета конвейерной линии", level=1)
        for key, value in data.items():
            doc.add_paragraph(f"{key}: {value}")
        try:
            doc.save(filename)
            messagebox.showinfo("Экспорт", f"Данные успешно экспортированы в {filename}")
        except Exception as e:
            messagebox.showerror("Ошибка экспорта", str(e))

    def export_to_xlsx(self, filename, data):
        try:
            workbook = openpyxl.load_workbook(filename)
            sheet = workbook.active
            # Находим первую пустую строку
            row = sheet.max_row + 1

        except FileNotFoundError:
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            # Заголовки
            sheet['A1'] = 'Параметр'
            sheet['B1'] = 'Значение'
            row = 2

        for key, value in data.items():
            sheet.cell(row=row, column=1, value=key)
            sheet.cell(row=row, column=2, value=value)
            row += 1

        try:
            workbook.save(filename)
            messagebox.showinfo("Экспорт", f"Данные успешно экспортированы в {filename}")
        except Exception as e:
            messagebox.showerror("Ошибка экспорта", str(e))

# View
""" Класс представления """
class MainWindow:
    # Конструктор класса
    def __init__(self, controller):
        self.controller = controller
        self.window = tk.Tk()

        # Регистрация команды валидирования float данных
        self.validate_float_cmd = self.window.register(self.validate_float)

        # Настройка окна
        self.window.title("Вариант 19. Определение требуемого количества постов конвейерной линии.")
        self.window.geometry("700x400")
        self.window.minsize(700, 400)

        # Настройка сетки
        self.window.columnconfigure(0, weight=1)
        self.window.columnconfigure(1, weight=1)

        self.window.rowconfigure(0, weight=1)
        self.window.rowconfigure(1, weight=1)
        
        # Создание виджетов ввода данных

        # Виджет выбора типа операции
        self.operation_type_label = ttk.Label(self.window, text="Вид операций:")
        self.operation_type_combo = ttk.Combobox(self.window, values=[e.value for e in OperationType], state="readonly")

        # Виджет выбора сложности изделия
        self.product_complexity_label = ttk.Label(self.window, text="Сложность изделия:")
        self.product_complexity_combo = ttk.Combobox(self.window, values=[e.value for e in ProductComplexity], state="readonly")

        # Виджет ввода средней продолжительности операции на доформовочном участке(мин):
        self.preform_operation_time_label = ttk.Label(self.window, text="Средняя продолжительность операции на доформовочном участке(мин):")
        # Валидация entry происходит при нажатии клавиши(validate="key"), %P - подстановка после применения изменений
        self.preform_operation_time_entry = ttk.Entry(self.window, validate="key", validatecommand=(self.validate_float_cmd, '%P'))

        # Виджет ввода средней продолжительности операции на формовочном участке(мин):
        self.form_operation_time_label = ttk.Label(self.window, text="Средняя продолжительность операции на формовочном участке(мин):")
        self.form_operation_time_entry = ttk.Entry(self.window, validate="key", validatecommand=(self.validate_float_cmd, '%P'))

        # Виджет ввода средней продолжительности операции на послеформовочном участке(мин):
        self.postform_operation_time_label = ttk.Label(self.window, text="Средняя продолжительность операции на послеформовочном участке(мин):")
        self.postform_operation_time_entry = ttk.Entry(self.window, validate="key", validatecommand=(self.validate_float_cmd, '%P'))

        # Виджет ввода продолжительности цикла формования
        self.cycle_time_label = ttk.Label(self.window, text="Продолжительность цикла формования (мин):")
        self.cycle_time_entry = ttk.Entry(self.window, validate="key", validatecommand=(self.validate_float_cmd, '%P'))

        # Переменная для хранения значения ползунка
        self.transfer_time = tk.DoubleVar(value=1.5)
        # Виджет ввода продолжительности передвижения тележек
        self.transfer_time_label = ttk.Label(self.window, text="Продолжительность передвижения тележек (мин):")

        self.transfer_time_entry = tk.Scale(
            self.window,
            variable=self.transfer_time,
            from_=MIN_TRANSFER_TIME,
            to=MAX_TRANSFER_TIME,
            resolution=0.1,
            orient="horizontal",
            showvalue=True
            )

        # Виджет ввода коэффициента
        self.denominator_var = tk.BooleanVar()
        self.denomirator_label = ttk.Label(self.window, text="Использовать знаменатель коэффициента?")
        self.denomirator_check = ttk.Checkbutton(self.window, variable=self.denominator_var)

        # Кнопка расчета
        self.calculate_button = ttk.Button(self.window, text="Рассчитать", command=self.calculate)

        # Виджет вывода результата
        self.result_label = ttk.Label(self.window, text="Результат:")
        self.result_value = tk.StringVar()
        self.result_display = ttk.Label(self.window, textvariable=self.result_value)

        # Кнопка экспорта
        self.doc_button = ttk.Button(self.window, text="Экспортировать *.docx")
        self.xls_button = ttk.Button(self.window, text="Экспортировать *.xlsx")

        # Размещение виджетов

        self.operation_type_label.grid(row=0, column=0, sticky=tk.EW, padx=5, pady=5)
        self.operation_type_combo.grid(row=0, column=1, sticky=tk.EW, padx=5, pady=5)

        self.product_complexity_label.grid(row=1, column=0, sticky=tk.EW, padx=5, pady=5)
        self.product_complexity_combo.grid(row=1, column=1, sticky=tk.EW, padx=5, pady=5)

        self.preform_operation_time_label.grid(row=2, column=0, sticky=tk.EW, padx=5, pady=5)
        self.preform_operation_time_entry.grid(row=2, column=1, sticky=tk.EW, padx=5, pady=5)

        self.form_operation_time_label.grid(row=3, column=0, sticky=tk.EW, padx=5, pady=5)
        self.form_operation_time_entry.grid(row=3, column=1, sticky=tk.EW, padx=5, pady=5)

        self.postform_operation_time_label.grid(row=4, column=0, sticky=tk.EW, padx=5, pady=5)
        self.postform_operation_time_entry.grid(row=4, column=1, sticky=tk.EW, padx=5, pady=5)

        self.cycle_time_label.grid(row=5, column=0, sticky=tk.EW, padx=5, pady=5)
        self.cycle_time_entry.grid(row=5, column=1, sticky=tk.EW, padx=5, pady=5)

        self.transfer_time_label.grid(row=6, column=0, sticky=tk.EW, padx=5, pady=5)
        self.transfer_time_entry.grid(row=6, column=1, sticky=tk.EW, padx=5, pady=5)

        self.denomirator_label.grid(row=7, column=0, sticky=tk.EW, padx=5, pady=5)
        self.denomirator_check.grid(row=7, column=1, sticky=tk.EW, padx=5, pady=5)

        self.calculate_button.grid(row=8, column=0, columnspan=2, sticky=tk.EW, padx=5, pady=5)

        self.result_label.grid(row=9, column=0, sticky=tk.EW, padx=5, pady=5)
        self.result_display.grid(row=9, column=1, sticky=tk.EW, padx=5, pady=5)

        self.doc_button.grid(row=10, column=0, columnspan=2, sticky=tk.EW, padx=5, pady=5)
        self.xls_button.grid(row=11, column=0, columnspan=2, sticky=tk.EW, padx=5, pady=5)

    """Метод валидации чисел с плавающей точкой"""
    def validate_float(self, new_value):
        if new_value == "":
            return True
        try:
            float(new_value)
            return True
        except ValueError:
            return False
        
    def calculate(self):
        try:
            # Извлечение значений из полей ввода
            operation_type = OperationType(self.operation_type_combo.get())
            product_complexity = ProductComplexity(self.product_complexity_combo.get())
            preform_operation_time = float(self.preform_operation_time_entry.get())
            form_operation_time = float(self.form_operation_time_entry.get())
            postform_operation_time = float(self.postform_operation_time_entry.get())
            cycle_time = float(self.cycle_time_entry.get())
            transfer_time = float(self.transfer_time_entry.get())
            use_denominator = bool(self.denominator_var.get())

            # Вызов метода расчета в контроллере
            self.controller.calculate_and_display(operation_type, product_complexity, preform_operation_time, form_operation_time, postform_operation_time, cycle_time, transfer_time, use_denominator)
        except ValueError as e:
            self.result_value.set("Ошибка: Введены некорректные данные.")
            messagebox.showerror("Ошибка", e)
        except ZeroDivisionError:
            self.result_value.set("Ошибка: В ходе расчета произошло деление на ноль. Проверьте корректность данных.")
        except Exception as e:
            self.result_value.set("Ошибка: Произошла непредвиденная ошибка.")
            messagebox.showerror("Ошибка", e)

    # Метод для заполнения поля результата
    def set_result(self, result):
        self.result_value.set(str(result))

    # Метод запуска главного цикла обработки событий окна
    def run(self):
        self.window.mainloop()

class CalculatorController:
    def __init__(self, view, calculator):
        # Текцущее представление
        self.view = view
        # Сервис для подсчёта количества постов
        self.calculator = calculator
        # Словарь коэффициентов
        self.coefficients = {
            (OperationType.AUTOMATED, ProductComplexity.SIMPLE): OperationCoefficients(OperationType.AUTOMATED, ProductComplexity.SIMPLE, 1.05),
            (OperationType.AUTOMATED, ProductComplexity.COMPLEX): OperationCoefficients(OperationType.AUTOMATED, ProductComplexity.COMPLEX, 1.05),
            (OperationType.MECHANIZED, ProductComplexity.SIMPLE): OperationCoefficients(OperationType.MECHANIZED, ProductComplexity.SIMPLE, 1.15, 1.10),
            (OperationType.MECHANIZED, ProductComplexity.COMPLEX): OperationCoefficients(OperationType.MECHANIZED, ProductComplexity.COMPLEX, 1.25, 1.15),
            (OperationType.MANUAL, ProductComplexity.SIMPLE): OperationCoefficients(OperationType.MANUAL, ProductComplexity.SIMPLE, 1.25, 1.15),
            (OperationType.MANUAL, ProductComplexity.COMPLEX): OperationCoefficients(OperationType.MANUAL, ProductComplexity.COMPLEX, 1.35, 1.20),
        }
    def calculate_and_display(self, operation_type, product_complexity, preform_operation_time, form_operation_time, postform_operation_time, cycle_time, transfer_time, use_denominator):
        segment = ConveyorLineSegment(operation_type, product_complexity, preform_operation_time, form_operation_time, postform_operation_time, cycle_time, transfer_time)

        # Поиск коэффициента для заданной пары значений
        coefficient = self.coefficients.get((operation_type, product_complexity))
        if not coefficient:
            self.view.set_result("Ошибка: Не найден коэффициент для указанных параметров")
            return

        # Вычисление количества постов
        result = self.calculator.calculate_posts(segment, coefficient, use_denominator)

        # Форматирование результата
        formatted_result = "{:.2f}".format(result)

        # Отображение результата
        self.view.set_result(formatted_result)

# Запуск приложения
if __name__ == "__main__":
    calculator = Calculator()
    view = MainWindow(None)
    controller = CalculatorController(view, calculator)
    view.controller = controller
    view.run()